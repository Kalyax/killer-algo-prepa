from random import choice, randint
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

path_dossier=Path.cwd().__str__() #"chemin d'accès du dossier algo_killer (exemple : C:/Documents/BDP/ :")

path_fichier_boucle_modele = path_dossier+"\\docs_vierges\\boucle_modele.docx"
path_fichier_boucle = path_dossier+"\\boucle.docx"
path_fichier_cartes_modele = path_dossier+"\\docs_vierges\\cartes_modele.docx"
path_fichier_cartes = path_dossier+"\\cartes.docx"
path_classes  = path_dossier+"\\classes\\"
print('indiquer le nom de chaque classe (le nom doit être celui du fichier texte sans l\'extension, exemple : PC). Entrer "/" lorsque toutes les classes ont été entrées')
classe_a_ajouter = input("classe 1 :")
c=1
classes=["MP","HK","KH","MPSI","MP2I","MPI","PC","PSI","PCSI","ECG1","ECG2"]
while classe_a_ajouter != "/" :
    c+=1
    classes.append(classe_a_ajouter)
    classe_a_ajouter = input("classe "+str(c)+" :")

print(path_classes)

def etudiants():
    etudiants = []
    for classe in classes:
        path_classe = path_classes + classe + ".txt"
        with  open(path_classe, 'r', encoding='utf-8') as fichier_classe:
            exec(classe + '=fichier_classe.read().splitlines()')
            exec("etudiants.append(1*" + classe + ")")
    for c in range(len(etudiants)):
        for e in range(len(etudiants[c])):
            etudiants[c][e] = [1 * etudiants[c][e], classes[c]]
    nbre_participants = 0
    for classe in etudiants:
        nbre_participants += len(classe)
    return etudiants


def nbre_joueurs(et) :
    n=0
    for classe in et :
        n+=len(classe)
    return n, str(len(et))+" classes et "+str(n)+" etudiants à placer dans la boucle"


def supprimer_elements_liste(liste, elements) :
    e = 0
    while e < len(liste) :
        if liste[e] in elements :
            liste.pop(e)
        else :
            e+=1
    return 1*liste


def supprimer_elements_liste_double(liste_double, elements) :
    for c in range(len(liste_double)) :
        supprimer_elements_liste(liste_double[c], elements)
    return 1*liste_double


def supprimer_elements_vides(liste) :
    return supprimer_elements_liste(liste, [[]])


def classemax(et) :
    max=len(et[0])
    ind=0
    for c in range(len(et)) :
        if len(et[c]) > max :
            max=len(et[c])
            ind=c
    return ind


def test(boucle, a_ajouter) :
    for j in range(1,4) :
        for j_ in range(len(a_ajouter)) :
            if boucle[-j][1]==a_ajouter[j_][1] or len(a_ajouter) == 0 :
                return False
    # print("trou2")
    return True


def ajouter_4(et) :
    ind_classe_max = classemax(et)
    joueur1 = choice(et[ind_classe_max])
    a_ajouter = [joueur1]
    completage = et[:ind_classe_max] + et[ind_classe_max + 1:]
    # print(completage, "1")
    completage=supprimer_elements_vides(completage)
    # print(completage, "2")
    for c in range(3):
        # print(len(completage))
        ind_class_ajout = randint(0, len(completage) - 1)
        classe_ajout = completage[ind_class_ajout]
        # print(classe_ajout)
        joueur_ajout = choice(classe_ajout)
        a_ajouter.append(joueur_ajout)
        completage.pop(ind_class_ajout)
    return a_ajouter


def inserable(boucle, j, etudiant) :
    if j < 4 :
        return False
    for j_ in range(3) :
        if boucle[j-j_-1][1] == etudiant[1] or boucle[j+j_][1] == etudiant[1] :
            return False
    return True

def boucle_killer() :
    et = 1 * etudiants()
    boucle = ajouter_4(et)
    et=supprimer_elements_vides(et)
    et=supprimer_elements_liste_double(et, boucle)
    a_ajouter = ajouter_4(et)
    while len(et) >= 4 :
        compteur_secu=0
        while not test(boucle, a_ajouter) :
            compteur_secu+=1
            if compteur_secu== 100 :
                return boucle_killer()
            # print("québlo")
            # print(len(boucle))
            a_ajouter=ajouter_4(et)
        boucle=boucle+a_ajouter
        # print(boucle)
        et = supprimer_elements_liste_double(et, a_ajouter)
        et = supprimer_elements_vides(et)
    # print(et, len(et))
    for c in range(len(et)) :
        for e in range(len(et[c])) :
            j=0
            while not inserable(boucle, j, et[c][e]) :
                j+=1
                if j == len(boucle)-3 :
                    return boucle_killer()
            boucle = boucle[:j]+[et[c][e]]+boucle[j:]
    return boucle


def verif_finale(boucle) :
    if len(boucle) < nbre_joueurs(etudiants())[0] :
        return False
    for i in range(len(boucle)) :
        for j in range(len(boucle)) :
            if boucle[i][0]==boucle[j][0] and i!= j :
                return False
    for j in range(len(boucle)) :
        for j_ in range(1,4) :
            if boucle[j-j_][1]==boucle[j][1] :
                return False
    return True


print("Récupération des données de chaque classe")
print(nbre_joueurs(etudiants())[1])
print("Recherche d'une boucle adéquate")
boucle_jeu=boucle_killer()
while not verif_finale(boucle_jeu)  :
    boucle_jeu=boucle_killer()
print("Boucle trouvée")


print("Création du fichier boucle ("+ path_fichier_boucle+")")
texte_fichier_boucle = ""
for joueur in boucle_jeu :
    texte_fichier_boucle = texte_fichier_boucle + joueur[0]+" (" + joueur[1] + ") -> "


fichier_boucle = Document(path_fichier_boucle_modele)
paragraphe = fichier_boucle.add_paragraph()
para = paragraphe.add_run()
font = para.font
font.name = 'Calibri'
font.size = Pt(11)
para.text = texte_fichier_boucle
fichier_boucle.save(path_fichier_boucle)


print("Création du fichier cartes ("+path_fichier_cartes+")")
cartes = []
for j in range(-1, len(boucle_jeu)-1) :
    if boucle_jeu[j][1] not in [cartes[k][0][0][1] for k in range(len(cartes))] :
        cartes.append([[boucle_jeu[j], boucle_jeu[j+1]]])
    else :
        for classe in cartes :
            if classe[0][0][1] == boucle_jeu[j][1] :
                classe.append([boucle_jeu[j], boucle_jeu[j+1]])




indices_cartes = []
for c in range(len(cartes)) :
    indices_cartes.append(cartes[c][0][0][1])
    for j in range(len(cartes[c])) :
        indices_cartes.append([c, j])
    for k in range(8-(len(cartes[c])+1)%8) :
        indices_cartes.append('')


def nbre_cartes(et) :
    s=0
    for classe in et :
        s+=len(classe)+1+(8-(len(classe)+1)%8)
    return s

indices_killers = [k for k in range(nbre_cartes(etudiants())*2) if k%16 <8]
indices_cibles = [k for k in range(nbre_cartes(etudiants())*2) if k%16 >=8]
a_imprimer = ['' for k in range(len(indices_killers)*2)]
for k in range(len(indices_killers)) :
    if k < len(indices_cartes) :
        a=indices_killers[k]
        if indices_cartes[k] != '' :
            if indices_cartes[k] not in classes :
                a_imprimer[indices_killers[k]] = cartes[indices_cartes[k][0]][indices_cartes[k][1]][0][0]
                a_imprimer[indices_cibles[k]] = "cible : " + cartes[indices_cartes[k][0]][indices_cartes[k][1]][1][0]
            else :
                a_imprimer[indices_killers[k]] = indices_cartes[k]



fichier_cartes = Document(path_fichier_cartes_modele)
tableaux = fichier_cartes.tables
cellules=[]
for k in range(len(tableaux)) :
    for i in range(4) :
        for j in range(2) :
            cellules.append([k, i, j])



for j in range(len(a_imprimer)) :
    if cellules[j][0]%2==0 :
        cellule = tableaux[cellules[j][0]].cell(cellules[j][1], cellules[j][2])
    else :
        if cellules[j][2] == 0 :
            cellule = tableaux[cellules[j][0]].cell(cellules[j][1], 1)
        else :
            cellule = tableaux[cellules[j][0]].cell(cellules[j][1], 0)
    paragraphe = cellule.add_paragraph()
    para = paragraphe.add_run()
    paragraphe.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font = para.font
    font.name = 'Modern Love'
    font.size = Pt(18)
    para.text = "\n \n " + a_imprimer[j]


fichier_cartes.save(path_fichier_cartes)
print("Programme terminé, cartes à imprimer ;)")